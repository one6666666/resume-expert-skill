#!/usr/bin/env python3
"""
Resume DOCX Generator — 生成 ATS 兼容的 Word 简历文档（2026 AI 时代版）

用法:
    python generate_docx.py <output_path> <json_data_path>

JSON 数据格式:
{
    "name": "张三",
    "target_position": "高级后端工程师（具备 AI Agent 全栈开发经验）",
    "years": "5",
    "email": "zhangsan@example.com",
    "phone": "13800138000",
    "location": "北京",
    "github": "zhangsan",
    "portfolio": "https://zhangsan.dev",
    "summary": "个人简介文本...",
    "skills": [
        {"name": "Java", "ai": false},
        {"name": "Spring AI", "ai": true},
        {"name": "RAG / MCP", "ai": true}
    ],
    "experience": [
        {
            "role": "高级Java开发工程师",
            "company": "某科技有限公司",
            "date": "2021.06 - 至今",
            "product_info": "核心支付平台，日活 50 万+",
            "details": [
                "主导设计高并发支付系统，日均处理交易100万+，可用性99.99%",
                "将大模型能力融入业务系统，违规内容拦截率提升至98%"
            ]
        }
    ],
    "projects": [
        {
            "name": "AI 超级智能体",
            "url": "https://agent.example.com",
            "tech": "Java, Spring AI, LangChain4j, RAG, MCP",
            "description": "从0到1搭建AI超级智能体平台，支持多模型接入..."
        }
    ],
    "other_works": [
        {"name": "轻量级RPC框架", "tech": "Java, Netty", "highlight": "自研RPC通信协议"}
    ],
    "education": {
        "degree": "计算机科学与技术 本科",
        "school": "某大学",
        "date": "2015.09 - 2019.06",
        "edu_note": "GPA 3.8/4.0，排名前10% | 校优秀毕业生"
    },
    "strengths": [
        "全栈能力强：已上线10+个可访问产品",
        "热爱开源：GitHub 2000+ Followers"
    ]
}

依赖:
    pip install python-docx
"""

import json
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── 颜色配置 ──
PRIMARY = RGBColor(0x00, 0x33, 0x66)   # 深蓝
AI_TAG_COLOR = RGBColor(0x1A, 0x5C, 0x1A)  # AI技能标签绿
TEXT = RGBColor(0x33, 0x33, 0x33)
TEXT_LIGHT = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT = RGBColor(0x66, 0x88, 0xAA)


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text):
    """添加深蓝色标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 底部边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return p


def add_body_text(doc, text, bold=False, size=10, color=TEXT):
    """添加正文文本"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    return p


def add_small_text(doc, text, size=9, color=TEXT_LIGHT):
    """添加小号辅助文本"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullet(doc, text, size=10):
    """添加项目符号条目"""
    p = doc.add_paragraph()
    prefix_run = p.add_run('• ')
    prefix_run.font.size = Pt(size)
    prefix_run.font.color.rgb = PRIMARY
    prefix_run.font.name = '微软雅黑'
    prefix_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_LIGHT
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_header_line(doc, text, bold=False, size=11):
    """添加头部行"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = WHITE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_header_link(doc, text, size=9):
    """添加头部链接行"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_skill_tags_inline(doc, skills, size=10):
    """添加技能标签行（AI技能绿色高亮）"""
    p = doc.add_paragraph()

    for i, skill in enumerate(skills):
        if isinstance(skill, str):
            name = skill
            is_ai = False
        else:
            name = skill.get('name', '')
            is_ai = skill.get('ai', False)

        if i > 0:
            sep = p.add_run('  ')
            sep.font.size = Pt(size)

        run = p.add_run(name)
        run.font.size = Pt(size)
        run.font.color.rgb = AI_TAG_COLOR if is_ai else TEXT
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        if i < len(skills) - 1:
            sep2 = p.add_run(' | ')
            sep2.font.size = Pt(size - 1)
            sep2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    p.paragraph_format.space_after = Pt(4)
    return p


def generate_resume(data, output_path):
    """生成简历 DOCX"""

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 头部区域 ──
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    set_cell_shading(header_cell, '003366')

    name = data.get('name', '姓名')
    target = data.get('target_position', '')
    years = data.get('years', '')

    p = header_cell.paragraphs[0]
    run = p.add_run(name)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle_text = f"{target}  |  {years}年经验" if target and years else target or years or ''
    if subtitle_text:
        add_header_line(header_cell, subtitle_text, size=9)

    # 联系方式
    contact_parts = []
    if data.get('email'):
        contact_parts.append(data['email'])
    if data.get('phone'):
        contact_parts.append(data['phone'])
    if data.get('location'):
        contact_parts.append(data['location'])

    if contact_parts:
        add_header_line(header_cell, ' | '.join(contact_parts), size=9)

    # 作品集 / GitHub / LinkedIn
    link_parts = []
    if data.get('github'):
        link_parts.append(f"github.com/{data['github']}")
    if data.get('portfolio'):
        link_parts.append(f"{data['portfolio']}")
    if data.get('linkedin'):
        link_parts.append(data['linkedin'])

    if link_parts:
        add_header_link(header_cell, ' | '.join(link_parts), size=9)

    doc.add_paragraph()  # 间距

    # ── 个人简介 ──
    summary = data.get('summary', '')
    if summary:
        add_heading_styled(doc, '个人简介')
        add_body_text(doc, summary, size=10, color=TEXT_LIGHT)
        doc.add_paragraph()

    # ── 专业技能 ──
    skills = data.get('skills', [])
    if skills:
        add_heading_styled(doc, '专业技能')
        add_skill_tags_inline(doc, skills, size=10)
        doc.add_paragraph()

    # ── 工作经历 ──
    experience = data.get('experience', [])
    if experience:
        add_heading_styled(doc, '工作经历')

        for exp in experience:
            role = exp.get('role', '')
            company = exp.get('company', '')
            date = exp.get('date', '')
            product_info = exp.get('product_info', '')

            header_text = f"{role}"
            if company:
                header_text += f" — {company}"
            if date:
                header_text += f"  ({date})"
            add_body_text(doc, header_text, bold=True, size=10, color=TEXT)

            if product_info:
                add_small_text(doc, f"📦 {product_info}", size=9, color=ACCENT)

            for detail in exp.get('details', []):
                add_bullet(doc, detail, size=10)

            doc.add_paragraph()

    # ── 项目经验 ──
    projects = data.get('projects', [])
    if projects:
        add_heading_styled(doc, '项目经验')

        for proj in projects:
            name = proj.get('name', '项目名称')
            url = proj.get('url', '')

            # 项目名 + 链接
            name_line = name
            if url:
                name_line += f"  [{url}]"
            add_body_text(doc, name_line, bold=True, size=10, color=PRIMARY)

            tech = proj.get('tech', '')
            if tech:
                add_small_text(doc, f"技术栈：{tech}", size=9, color=PRIMARY)

            desc = proj.get('description', '')
            if desc:
                add_bullet(doc, desc, size=10)

            doc.add_paragraph()

    # ── 其他个人作品 ──
    other_works = data.get('other_works', [])
    if other_works:
        add_heading_styled(doc, '其他个人作品')

        portfolio = data.get('portfolio', '')
        if portfolio:
            add_small_text(doc, f"更多项目详见 → {portfolio}", size=9, color=ACCENT)

        for work in other_works:
            name = work.get('name', '')
            tech = work.get('tech', '')
            highlight = work.get('highlight', '')
            text = f"• {name}"
            if tech:
                text += f" ({tech})"
            if highlight:
                text += f"：{highlight}"
            add_small_text(doc, text, size=10, color=TEXT_LIGHT)

        doc.add_paragraph()

    # ── 教育背景 ──
    edu = data.get('education', {})
    if edu:
        add_heading_styled(doc, '教育背景')
        degree = edu.get('degree', '')
        school = edu.get('school', '')
        date = edu.get('date', '')
        edu_text = f"{degree} — {school}"
        if date:
            edu_text += f"  ({date})"
        add_body_text(doc, edu_text, size=10)

        edu_note = edu.get('edu_note', '')
        if edu_note:
            add_bullet(doc, edu_note, size=10)

        doc.add_paragraph()

    # ── 个人优势 ──
    strengths = data.get('strengths', [])
    if strengths:
        add_heading_styled(doc, '个人优势')
        for s in strengths:
            add_bullet(doc, s, size=10)

    # 保存
    doc.save(output_path)
    print(f"✅ 简历已生成: {output_path}")
    return output_path


def main():
    if len(sys.argv) < 3:
        print("用法: python generate_docx.py <output_path.docx> <json_data_path.json>")
        print("  or: python generate_docx.py <output_path.docx>")
        print("       (从 stdin 读取 JSON)")
        sys.exit(1)

    output_path = sys.argv[1]

    if len(sys.argv) >= 3:
        json_path = sys.argv[2]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)

    generate_resume(data, output_path)


if __name__ == '__main__':
    main()
